import os
import uuid
import datetime
import shutil
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, BackgroundTasks
from sqlalchemy.orm import Session
from typing import List

from app.database.connection import get_db
from app.models.models import User, KnowledgeDocument
from app.schemas.schemas import ApiResponse, KnowledgeDocumentBase, KnowledgeStats
from app.routers.auth import get_current_user
from app.services.rag_service import RagService

router = APIRouter(prefix="/api/v1/admin", tags=["Admin Knowledge Base"])
rag_service = RagService()

UPLOAD_DIR = "public/uploads/kb"

# Ensure upload directory exists
os.makedirs(UPLOAD_DIR, exist_ok=True)

def verify_admin(current_user: User):
    if current_user.role not in ["super_admin", "ADMIN"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied. Administrators only."
        )

@router.post("/upload-document", response_model=ApiResponse[KnowledgeDocumentBase])
def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    category: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    verify_admin(current_user)
    
    # 1. Validate file extension
    filename = file.filename or "unknown"
    ext = filename.split(".")[-1].lower() if "." in filename else ""
    if ext not in ["pdf", "docx", "txt", "md"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file format. Only PDF, DOCX, TXT, and Markdown files are allowed."
        )
        
    # 2. Save file locally
    doc_id = str(uuid.uuid4())
    safe_filename = f"{doc_id}_{filename}"
    file_path = os.path.join(UPLOAD_DIR, safe_filename)
    
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to write file to disk: {e}"
        )
        
    # 3. Create document record
    now_str = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    new_doc = KnowledgeDocument(
        id=doc_id,
        filename=filename,
        category=category,
        file_type=ext,
        upload_date=now_str,
        status="Processing",
        chunk_count=0,
        indexed_status=False,
        file_path=file_path,
        created_at=now_str,
        updated_at=now_str
    )
    
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    
    # 4. Trigger background task
    background_tasks.add_task(
        rag_service.process_and_index_document,
        doc_id=doc_id,
        file_path=file_path,
        filename=filename,
        category=category,
        file_type=ext,
        db=None
    )
    
    return ApiResponse(data=new_doc)

@router.get("/documents", response_model=ApiResponse[List[KnowledgeDocumentBase]])
def get_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    verify_admin(current_user)
    docs = db.query(KnowledgeDocument).order_by(KnowledgeDocument.upload_date.desc()).all()
    return ApiResponse(data=docs)

@router.delete("/document/{id}", response_model=ApiResponse[dict])
def delete_document(
    id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    verify_admin(current_user)
    
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == id).first()
    if not doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found."
        )
        
    # Delete from Pinecone
    rag_service.delete_document_vectors(id)
    
    # Delete local file
    if doc.file_path and os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except Exception as e:
            print(f"[RAG-Index] Error deleting file {doc.file_path}: {e}")
            
    # Delete from database
    db.delete(doc)
    db.commit()
    
    return ApiResponse(data={"success": True, "message": "Document deleted successfully."})

@router.post("/reindex/{id}", response_model=ApiResponse[KnowledgeDocumentBase])
def reindex_document(
    id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    verify_admin(current_user)
    
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == id).first()
    if not doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found."
        )
        
    # Verify file still exists on disk
    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Original document file not found on disk. Cannot reindex."
        )
        
    # Set status back to Processing
    doc.status = "Processing"
    doc.indexed_status = False
    doc.updated_at = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    db.commit()
    db.refresh(doc)
    
    # Clean up existing vectors in Pinecone before reindexing
    rag_service.delete_document_vectors(id)
    
    # Trigger background task
    background_tasks.add_task(
        rag_service.process_and_index_document,
        doc_id=doc.id,
        file_path=doc.file_path,
        filename=doc.filename,
        category=doc.category,
        file_type=doc.file_type,
        db=None
    )
    
    return ApiResponse(data=doc)

@router.get("/statistics", response_model=ApiResponse[KnowledgeStats])
def get_statistics(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    verify_admin(current_user)
    
    # Count total documents
    total_docs = db.query(KnowledgeDocument).count()
    
    # Sum chunk count
    total_chunks = 0
    docs = db.query(KnowledgeDocument).all()
    for d in docs:
        if d.chunk_count:
            total_chunks += d.chunk_count
            
    # Last updated
    latest_indexed = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.status == "Indexed"
    ).order_by(KnowledgeDocument.updated_at.desc()).first()
    
    last_updated = latest_indexed.updated_at if latest_indexed else None
    
    # Connectivity status checks
    pinecone_status = "Unavailable"
    if rag_service.pinecone_key:
        try:
            from pinecone import Pinecone
            pc = Pinecone(api_key=rag_service.pinecone_key)
            pc.list_indexes()
            pinecone_status = "Operational"
        except Exception:
            pass
            
    groq_status = "Operational" if rag_service.groq_key else "Unavailable"
    
    stats = KnowledgeStats(
        total_documents=total_docs,
        total_chunks=total_chunks,
        total_embeddings=total_chunks, # 1:1 mapping
        last_updated=last_updated,
        pinecone_status=pinecone_status,
        groq_status=groq_status
    )
    
    return ApiResponse(data=stats)
    
@router.get("/document/{id}/chunks", response_model=ApiResponse[List[str]])
def get_document_chunks(
    id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    verify_admin(current_user)
    
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == id).first()
    if not doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found."
        )
        
    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Original document file not found on disk. Cannot read chunks."
        )
        
    try:
        text = ""
        file_path = doc.file_path
        file_type = doc.file_type
        
        if file_type == "pdf":
            import pypdf
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        elif file_type == "docx":
            import docx
            docx_doc = docx.Document(file_path)
            text = "\n".join([para.text for para in docx_doc.paragraphs])
        else: # txt or md
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
                
        chunks = rag_service._split_text_recursive(text, max_chunk_size=1000, overlap=200)
        return ApiResponse(data=chunks)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to read chunks: {str(e)}"
        )

@router.get("/search-history", response_model=ApiResponse[List[dict]])
def get_search_history(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    verify_admin(current_user)
    
    from app.models.models import SearchHistory
    logs = db.query(SearchHistory).order_by(SearchHistory.id.desc()).limit(100).all()
    data = []
    for l in logs:
        data.append({
            "id": l.id,
            "query": l.query,
            "response": l.response,
            "timestamp": l.timestamp
        })
    return ApiResponse(data=data)

@router.post("/rag-playground", response_model=ApiResponse[dict])
def rag_playground(
    payload: dict,
    current_user: User = Depends(get_current_user)
):
    """
    Admin RAG Playground — Test queries against the knowledge base.
    Returns retrieved chunks, similarity scores, sources, and generated answer.
    """
    verify_admin(current_user)
    
    question = payload.get("question", "").strip()
    if not question:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Question is required."
        )
    
    result = rag_service.rag_playground(question)
    return ApiResponse(data=result)


@router.get("/health/pinecone", response_model=ApiResponse[dict])
def health_check_pinecone(current_user: User = Depends(get_current_user)):
    verify_admin(current_user)
    try:
        if not rag_service.pinecone_key:
            raise ValueError("Pinecone API Key is not set in environment.")
            
        from pinecone import Pinecone
        pc = Pinecone(api_key=rag_service.pinecone_key)
        
        # Describe index
        desc = pc.describe_index(rag_service.pinecone_index)
        index_name = desc.name
        dimension = desc.dimension
        
        # Test vector loop: insert -> query -> delete
        index = pc.Index(rag_service.pinecone_index)
        
        import random
        test_id = f"health_check_{random.randint(1000, 9999)}"
        test_values = [random.uniform(-1.0, 1.0) for _ in range(dimension)]
        
        # Insert
        index.upsert(vectors=[{
            "id": test_id,
            "values": test_values,
            "metadata": {
                "document_id": "health_check_temp",
                "filename": "health_check_temp",
                "text": "Pinecone database health check probe vector.",
                "upload_date": datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
            }
        }])
        
        # Query
        q_res = index.query(vector=test_values, top_k=1, include_metadata=True)
        
        # Delete
        index.delete(ids=[test_id])
        
        # Fetch stats
        stats = index.describe_index_stats()
        total_vectors = stats.get("total_vector_count", 0)
        
        return ApiResponse(data={
            "status": "connected",
            "index": index_name,
            "vectors": total_vectors,
            "dimension": dimension,
            "metric": desc.metric,
            "message": "Vector insert, query, and cleanup loop completed successfully."
        })
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Pinecone health check failed: {str(e)}"
        )


@router.get("/health/groq", response_model=ApiResponse[dict])
def health_check_groq(current_user: User = Depends(get_current_user)):
    verify_admin(current_user)
    try:
        if not rag_service.groq_key:
            raise ValueError("Groq API Key is not set in environment.")
            
        if not rag_service.groq_client:
            raise ValueError("Groq Client is not initialized.")
            
        import time
        start_time = time.time()
        
        completion = rag_service.groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a health probe helper."},
                {"role": "user", "content": "Explain CampusConnect AI briefly in one sentence."}
            ],
            max_tokens=50,
            temperature=0.1
        )
        
        latency = round(time.time() - start_time, 2)
        response_text = completion.choices[0].message.content
        
        return ApiResponse(data={
            "status": "connected",
            "model": "llama-3.3-70b-versatile",
            "response_time": latency,
            "test_response": response_text,
            "message": "Groq API is operational and responded successfully."
        })
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Groq health check failed: {str(e)}"
        )


@router.get("/health/embedding", response_model=ApiResponse[dict])
def health_check_embedding(current_user: User = Depends(get_current_user)):
    verify_admin(current_user)
    try:
        if not rag_service.pinecone_key:
            raise ValueError("Pinecone API Key is not set in environment.")
            
        from pinecone import Pinecone
        pc = Pinecone(api_key=rag_service.pinecone_key)
        desc = pc.describe_index(rag_service.pinecone_index)
        dimension = desc.dimension
        
        # Test generation
        if dimension == 1024:
            embeddings_res = pc.inference.embed(
                model="multilingual-e5-large",
                inputs=["CampusConnect AI health probe text"],
                parameters={"input_type": "query", "truncate": "END"}
            )
            v_len = len(embeddings_res[0].values)
            model_name = "multilingual-e5-large"
        else: # 1536
            if not rag_service.openai_key:
                raise ValueError("Pinecone index dimension is 1536, but OPENAI_API_KEY is not set.")
            from openai import OpenAI
            client = OpenAI(api_key=rag_service.openai_key)
            res = client.embeddings.create(
                input=["CampusConnect AI health probe text"],
                model="text-embedding-3-small"
            )
            v_len = len(res.data[0].embedding)
            model_name = "text-embedding-3-small"
            
        return ApiResponse(data={
            "status": "connected",
            "model": model_name,
            "dimension": v_len,
            "message": f"Embedding generation successful using {model_name}."
        })
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Embedding health check failed: {str(e)}"
        )
